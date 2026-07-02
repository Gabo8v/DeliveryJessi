from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Colors ───
ORANGE = RGBColor(0xE8, 0x89, 0x5F)
ORANGE_DK = RGBColor(0xD4, 0x78, 0x50)
SAGE = RGBColor(0x7A, 0x9E, 0x7F)
SAGE_DK = RGBColor(0x6A, 0x8E, 0x6F)
BROWN = RGBColor(0x3D, 0x2C, 0x2E)
MUTED = RGBColor(0x8C, 0x7A, 0x7B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BROWN
style.paragraph_format.space_after = Pt(6)

# ─── Helper ───
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ORANGE_DK if level == 1 else BROWN
    return h

# ─── PORTADA ───
if os.path.exists('Logo.jpg'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('Logo.jpg', width=Cm(4), height=Cm(4))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comida Mecha')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BROWN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Gestión para Viandas Caseras')
run.font.size = Pt(14)
run.font.color.rgb = MUTED

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Presentación y Manual de Usuario')
run.font.size = Pt(14)
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.3.0')
run.font.size = Pt(12)
run.font.color.rgb = SAGE
run.font.italic = True

doc.add_page_break()

# ─── ÍNDICE ───
add_heading_styled('Índice', 1)
index_items = [
    ('1. Presentación', 0),
    ('2. Primeros pasos', 0),
    ('2.1 Pantalla principal', 1),
    ('2.2 Cargar el menú del día', 1),
    ('3. Gestión de pedidos', 0),
    ('3.1 Nuevo pedido', 1),
    ('3.2 Historial de pedidos', 1),
    ('4. Gestión de clientes', 0),
    ('4.1 Agregar cliente', 1),
    ('4.2 Cobrar deuda', 1),
    ('5. Reportes', 0),
    ('6. Ajustes', 0),
    ('6.1 Editar platos', 1),
    ('6.2 Menú del día', 1),
    ('6.3 Precios', 1),
    ('6.4 Actualizaciones', 1),
]
for text, indent in index_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent * 1)
    run = p.add_run(text)
    run.font.color.rgb = BROWN
    run.font.size = Pt(11)

doc.add_page_break()

# ─── 1. PRESENTACIÓN ───
add_heading_styled('1. Presentación', 1)
doc.add_paragraph(
    'Comida Mecha es una aplicación pensada para emprendimientos de comidas caseras. '
    'Nació de la necesidad de organizar los pedidos del día a día de forma simple, '
    'rápida y sin depender de conexión a internet.'
)
doc.add_paragraph(
    'Con Comida Mecha podés definir un menú diario de hasta cinco platos, tomar pedidos '
    'con todo lujo de detalles (tipo de entrega, método de pago, promociones), '
    'llevar el control de clientes fiados, registrar pagos, ver reportes diarios, '
    'semanales y mensuales, y hasta exportar los datos a Excel.'
)

# Info table
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
data = [
    ['Desarrollado para', 'Leonor y su emprendimiento de viandas'],
    ['Plataforma', 'Web (PWA) y Android'],
    ['Almacenamiento', 'Local (offline, no requiere internet)'],
    ['Versión actual', '1.3.0'],
]
for i, (k, v) in enumerate(data):
    cell0 = table.rows[i].cells[0]
    cell1 = table.rows[i].cells[1]
    cell0.text = k
    cell1.text = v
    for p in cell0.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = BROWN
            r.font.size = Pt(11)
    for p in cell1.paragraphs:
        for r in p.runs:
            r.font.color.rgb = BROWN
            r.font.size = Pt(11)
    set_cell_shading(cell0, 'FFF0E0')

doc.add_paragraph()
doc.add_paragraph(
    'La aplicación está diseñada para ser intuitiva: los botones principales están '
    'siempre visibles en la parte inferior de la pantalla, y cada acción importante '
    'tiene una confirmación para evitar errores.'
)

doc.add_page_break()

# ─── 2. PRIMEROS PASOS ───
add_heading_styled('2. Primeros pasos', 1)

add_heading_styled('2.1 Pantalla principal', 2)
doc.add_paragraph(
    'Al abrir la aplicación te encontrás con la pantalla de inicio. '
    'Allí te saluda Leonor (o el nombre que hayas configurado) y ves la fecha actual. '
    'Si ya cargaste el menú del día, los platos aparecen como tarjetas; '
    'tocando cualquiera de ellos vas directo a armar un pedido con ese plato preseleccionado.'
)
doc.add_paragraph(
    'En la parte inferior tenés cinco botones de navegación que te llevan a '
    'las secciones principales:'
)
nav_items = [
    '🏠 Inicio — Pantalla principal con el menú del día',
    '📋 Nuevo pedido — Para tomar un pedido nuevo',
    '📅 Pedidos — Historial de pedidos por día',
    '👥 Clientes — Lista de clientes y control de deudas',
    '⚙️ Ajustes — Configuración de platos, precios y más',
]
for item in nav_items:
    p = doc.add_paragraph(item, style='List Bullet')

add_heading_styled('2.2 Cargar el menú del día', 2)
doc.add_paragraph(
    'Antes de tomar pedidos, tenés que definir qué platos se preparan hoy. '
    'En la pantalla de inicio toca el botón "Cargar menú". '
    'Se abre una ventana donde podés buscar entre todos los platos disponibles '
    'y seleccionar hasta cinco. Cada plato seleccionado aparece como una etiqueta '
    'que podés sacar tocando la "x". Cuando termines, toca "Guardar".'
)
doc.add_paragraph(
    'Si querés cambiar el menú durante el día, volvé a tocar "Cargar menú" '
    'y ajustá los platos. Los pedidos ya registrados no se modifican.'
)

doc.add_page_break()

# ─── 3. GESTIÓN DE PEDIDOS ───
add_heading_styled('3. Gestión de pedidos', 1)

add_heading_styled('3.1 Nuevo pedido', 2)
doc.add_paragraph(
    'En la sección "Nuevo pedido" vas a armar cada pedido paso a paso:'
)
steps = [
    'Operador: escribí quién toma el pedido (por defecto aparece "Leonor").',
    'Platos: se muestran los platos del menú de hoy. Cada uno tiene un botón "Agregar" y control de cantidad. Si un plato tiene promoción (ej. 2 por $5000), aparece un botón especial.',
    'Cliente: escribí el nombre para buscar entre los clientes registrados. Si no existe, se habilita un formulario para crearlo al instante.',
    'Entrega: elegí "Retira" (pasa a buscar) o "A domicilio" (con costo de envío).',
    'Pago: elegí entre Efectivo, Transferencia/Mercado Pago, o Fiado. Si el cliente es fiado, aparece un aviso.',
    'Estado: podés marcarlo como "Pagado" directo (salvo que sea fiado, que queda pendiente).',
    'Tocá "Guardar pedido" y se almacena automáticamente.',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Consejo: ')
run.font.bold = True
run.font.color.rgb = ORANGE
p.add_run(
    'Si un cliente no aparece en la búsqueda, no hace falta '
    'salir de la pantalla. Escribí su nombre, tocá "Crear cliente" y completá '
    'los datos. Todo en el mismo lugar.'
)

add_heading_styled('3.2 Historial de pedidos', 2)
doc.add_paragraph(
    'En "Pedidos" podés navegar entre días con las flechas ‹ y ›. '
    'Se ven las viandas totales, lo cobrado y lo pendiente. '
    'Cada pedido muestra el cliente, los platos, el medio de pago, el tipo de entrega '
    'y el operador.'
)
doc.add_paragraph(
    'Desde acá podés marcar un pedido como "Pagado" cuando el cliente abona, '
    'o "Anular" si se cancela. Si el pedido era fiado, al pagarlo se descuenta '
    'automáticamente de la deuda del cliente.'
)

doc.add_page_break()

# ─── 4. GESTIÓN DE CLIENTES ───
add_heading_styled('4. Gestión de clientes', 1)

add_heading_styled('4.1 Agregar cliente', 2)
doc.add_paragraph(
    'En la pantalla "Clientes" podés ver todos los clientes registrados. '
    'Usá el buscador para filtrar por nombre o teléfono. '
    'Cada tarjeta muestra el nombre, teléfono y un indicador de deuda: '
    'rojo si debe plata, verde si está al día.'
)
doc.add_paragraph(
    'Tocá el botón "+" para agregar un cliente nuevo. '
    'Se abre una ventana con los campos: nombre completo, teléfono, dirección '
    'y un indicador para marcar si es cliente fiado.'
)

add_heading_styled('4.2 Cobrar deuda', 2)
doc.add_paragraph(
    'Tocando un cliente existente se abre un detalle con su información '
    'y el historial de los últimos pedidos. Si tiene deuda, acá podés '
    'registrar un pago: ingresá el monto y la fecha, y la deuda se actualiza sola. '
    'También podés ver todos los pagos registrados.'
)

doc.add_page_break()

# ─── 5. REPORTES ───
add_heading_styled('5. Reportes', 1)
doc.add_paragraph(
    'La sección "Reportes" te da una vista general del negocio. '
    'Podés elegir entre tres períodos:'
)
report_items = [
    'Día: muestra las estadísticas del día de hoy.',
    'Semana: resume los últimos siete días con un gráfico de barras que muestra la recaudación de cada día.',
    'Mes: estadísticas del mes actual.',
]
for item in report_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'En cada período ves: total de viandas vendidas, dinero total, plato más vendido, '
    'operador que más pedidos tomó y deuda total fiada. '
    'También hay una pestaña "Cierres" con el historial de cierres de turno.'
)
doc.add_paragraph(
    'Compartir: tocá el botón de compartir para copiar un resumen en texto '
    'y mandarlo por WhatsApp o donde quieras.'
)
doc.add_paragraph(
    'Descargar Excel: exportá los datos a un archivo de Excel '
    'para llevar tu propia contabilidad.'
)

doc.add_page_break()

# ─── 6. AJUSTES ───
add_heading_styled('6. Ajustes', 1)

add_heading_styled('6.1 Editar platos', 2)
doc.add_paragraph(
    'En "Ajustes" → "Editar Platos" ves la lista completa de platos '
    '(activos e inactivos). Tocando un plato se abre un formulario para modificar:'
)
plate_items = [
    'Nombre del plato',
    'Precio base',
    'Sopa (precio adicional si aplica)',
    'Promociones (ej. "2x5000", "3x7000")',
]
for item in plate_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph(
    'También podés activar o desactivar un plato con un interruptor. '
    'Los platos inactivos no aparecen en el menú del día.'
)

add_heading_styled('6.2 Menú del día', 2)
doc.add_paragraph(
    'También desde "Ajustes" podés cargar el menú para cualquier fecha, '
    'no solo la de hoy. Útil si querés planificar con anticipación.'
)

add_heading_styled('6.3 Precios', 2)
doc.add_paragraph(
    'En la sección de precios configurás el precio base de las viandas '
    'y el costo de envío a domicilio. Estos valores se usan como referencia '
    'al tomar pedidos.'
)

add_heading_styled('6.4 Actualizaciones', 2)
doc.add_paragraph(
    'Comida Mecha se actualiza periódicamente. En "Ajustes" → "Actualizaciones" '
    'podés verificar si hay una versión nueva disponible. Si la hay, '
    'podés descargar el APK directamente desde la aplicación.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—')
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Gracias por usar Comida Mecha.')
run.font.italic = True
run.font.color.rgb = MUTED

OUT = 'Presentacion_Comida_Mecha.docx'
doc.save(OUT)
print(f'DOCX generado: {OUT}')
